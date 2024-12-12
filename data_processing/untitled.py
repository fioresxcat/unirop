import os
import pdb
from pathlib import Path
import cv2
import numpy as np
import json
from omegaconf import OmegaConf
import fitz
import xml.etree.ElementTree as ET
from easydict import EasyDict
from docx import Document
from docx.shared import RGBColor
import re


def get_docx_reading_order():
    fp = 'raw_data/docx_files/de_kiem_tra.docx'

    document = Document(fp)
    words = []

    # Extract words from paragraphs
    for para in document.paragraphs:
        words.extend(para.text.split())

    # Extract words from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                words.extend(cell.text.split())

    with open('data_processing/test.txt', 'w') as f:
        for word in words:
            f.write(word + '\n')

    print('DONE')



def change_text_color_to_green(input_path, output_path):
    # Load the existing document
    document = Document(input_path)
    
    # Define green color (RGB: 0, 128, 0)
    green_color = RGBColor(0, 128, 0)
    
    # Iterate through all paragraphs and their runs to change text color
    for para in document.paragraphs:
        for run in para.runs:
            run.font.color.rgb = green_color

    # Iterate through all tables and their cells to change text color
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = green_color

    # Save to a new document
    document.save(output_path)


def color_words_in_sequence(input_path, output_path):
    # Load the existing document
    document = Document(input_path)

    # Define colors
    red = RGBColor(255, 0, 0)
    green = RGBColor(0, 128, 0)
    blue = RGBColor(0, 0, 255)
    colors = [red, green, blue]

    # Word counter to track the position of each word
    word_counter = 0

    # Function to split a run's text into words and retain formatting
    def colorize_run(paragraph, run):
        nonlocal word_counter
        words = re.findall(r'\S+|\s+', run.text)  # Split while keeping spaces
        
        # Keep track of the current run's formatting
        formatting = {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_size': run.font.size
        }

        # Clear the original run's text
        run.text = ''

        # Add new runs with appropriate colors
        for word in words:
            new_run = paragraph.add_run(word)
            new_run.bold = formatting['bold']
            new_run.italic = formatting['italic']
            new_run.underline = formatting['underline']
            new_run.font.size = formatting['font_size']
            new_run.font.color.rgb = colors[word_counter % 3]
            word_counter += 1

    # Iterate through all paragraphs and colorize each word
    for para in document.paragraphs:
        for run in para.runs:
            colorize_run(para, run)

    # Iterate through all tables and their cells to colorize words
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        colorize_run(para, run)

    # Save to a new document
    document.save(output_path)




if __name__ == '__main__':
    # get_docx_reading_order()
    # change_text_color_to_green(
    #     'raw_data/docx_files/de_kiem_tra.docx',
    #     'raw_data/docx_files/de_kiem_tra_green.docx'
    # )

    color_words_in_sequence(
        'raw_data/docx_files/de_kiem_tra.docx',
        'raw_data/docx_files/de_kiem_tra_colored.docx'
    )